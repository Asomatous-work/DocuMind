"""
Document Extractor — Text extraction for DOCX, DOC, and PDF files.
Bypasses OCR where possible; falls back to Vision OCR for scanned PDFs.
"""

import io
import time
import logging
from typing import Optional
import os

logger = logging.getLogger(__name__)

# MIME type and extension maps used by main.py
DOCUMENT_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "application/pdf": "pdf",
}

DOCUMENT_EXTENSIONS = {
    ".docx": "docx",
    ".doc": "doc",
    ".pdf": "pdf",
}


def _make_result(text: str, source_type: str, processing_time: float, confidence: float = 1.0) -> dict:
    """
    Build a result dict that matches the OCREngine output format.
    """
    # Split into pseudo-blocks (one per non-empty line/paragraph)
    raw_lines = [line.strip() for line in text.splitlines() if line.strip()]
    blocks = [
        {
            "text": line,
            "confidence": confidence,
            "bbox": [],
        }
        for line in raw_lines
    ]

    return {
        "text": text,
        "blocks": blocks,
        "block_count": len(blocks),
        "avg_confidence": confidence,
        "processing_time_seconds": round(processing_time, 2),
        "source_type": source_type,
    }


def extract_text_from_docx(file_bytes: bytes) -> dict:
    """Extract text from a .docx file."""
    start = time.time()
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is not installed.")

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    processing_time = time.time() - start
    return _make_result(full_text, "docx", processing_time)


def extract_text_from_pdf(file_bytes: bytes) -> dict:
    """Extract text from a PDF, with OCR fallback for scanned pages."""
    start = time.time()
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF is not installed.")

    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = []
    is_scanned = False

    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text("text").strip()
        if text:
            pages_text.append(f"--- Page {page_num} ---\n{text}")
        else:
            # Scanned page detected?
            is_scanned = True
            break # We'll handle full OCR logic if even one page is empty for simplicity or optimization
    
    if not is_scanned and pages_text:
        pdf.close()
        full_text = "\n\n".join(pages_text)
        return _make_result(full_text, "pdf", time.time() - start)

    # Fallback to OCR for scanned PDFs
    logger.info("⚠️ PDF appears to be scanned. Falling back to Vision OCR...")
    from ocr.engine import OCREngine
    engine = OCREngine()
    
    ocr_texts = []
    # Only OCR first 5 pages to avoid massive delays
    max_pages = min(5, len(pdf))
    
    for i in range(max_pages):
        page = pdf.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # Higher res for OCR
        img_bytes = pix.tobytes("png")
        
        result = engine.extract_text(img_bytes, source_type="pdf_scanned")
        ocr_texts.append(f"--- Page {i+1} (OCR) ---\n{result['text']}")
    
    pdf.close()
    full_text = "\n\n".join(ocr_texts)
    if len(pdf) > max_pages:
        full_text += f"\n\n... (Truncated: Only first {max_pages} pages processed via OCR)"
        
    return _make_result(full_text, "pdf_ocr", time.time() - start, confidence=0.9)


def extract_document(file_bytes: bytes, filename: str, mime_type: str) -> dict:
    """Route to the correct extractor."""
    ext = os.path.splitext(filename.lower())[1]
    
    # MIME check
    if mime_type == "application/pdf" or ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or ext == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported document type: {filename}")
